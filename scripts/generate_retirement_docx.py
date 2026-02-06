#!/usr/bin/env python3
"""
Comprehensive Retirement Planning Word Document Generator
Generates professional, editable retirement reports with all features.
"""

import json
import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def format_currency(amount):
    """Format currency in Australian style"""
    return f"${amount:,.0f}"

def format_percent(value):
    """Format percentage"""
    return f"{value:.1f}%"

def get_risk_level(success_rate):
    """Get risk level text and color"""
    if success_rate >= 85:
        return "LOW RISK", "22C55E"  # Green
    elif success_rate >= 70:
        return "MODERATE RISK", "F59E0B"  # Orange
    else:
        return "HIGH RISK", "EF4444"  # Red

def set_cell_background(cell, color_hex):
    """Set table cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_elm = OxmlElement(f'w:{edge}')
            edge_elm.set(qn('w:val'), 'single')
            edge_elm.set(qn('w:sz'), str(kwargs[edge]))
            edge_elm.set(qn('w:color'), kwargs.get('color', '000000'))
            tcBorders.append(edge_elm)
    tcPr.append(tcBorders)

def create_cover_page(doc, data):
    """Create professional cover page"""
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Australian Retirement Planning Report")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138)  # Navy blue
    
    doc.add_paragraph()  # Spacer
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Comprehensive Financial Analysis & Projections")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(71, 85, 105)  # Slate
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Key Info Table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    info = [
        ('Prepared For:', data.get('pensionRecipientType', 'single').title() + ' Retiree'),
        ('Current Age:', str(data.get('currentAge', 'N/A'))),
        ('Retirement Age:', str(data.get('retirementAge', 'N/A'))),
        ('Report Date:', datetime.now().strftime('%d %B %Y')),
    ]
    
    for i, (label, value) in enumerate(info):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        # Bold labels
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        # Color values
        table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Disclaimer
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = disclaimer.add_run("Important Disclaimer: ")
    run.font.bold = True
    run.font.size = Pt(9)
    run = disclaimer.add_run(
        "This report is for illustrative purposes only and does not constitute financial advice. "
        "The projections are based on assumptions that may not reflect actual future outcomes. "
        "Consult with a qualified financial adviser before making retirement decisions. "
        "Past performance is not indicative of future results."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_page_break()

def create_table_of_contents(doc):
    """Create table of contents"""
    
    heading = doc.add_heading("Table of Contents", level=1)
    heading.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    toc_items = [
        ("1. Executive Summary", "Key findings and portfolio health"),
        ("2. Financial Assumptions", "Input parameters and methodology"),
        ("3. Portfolio Projections", "Year-by-year balance and income"),
        ("4. Risk Analysis", "Monte Carlo results and stress testing"),
        ("5. Recommendations", "Actionable insights for your plan"),
        ("6. Scenario Details", "Detailed breakdowns and assumptions"),
    ]
    
    for title, desc in toc_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(12)
        
        p2 = doc.add_paragraph(desc)
        p2.paragraph_format.left_indent = Inches(0.5)
        p2.runs[0].font.size = Pt(10)
        p2.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_page_break()

def create_executive_summary(doc, data):
    """Create executive summary section"""
    
    heading = doc.add_heading("1. Executive Summary", level=1)
    heading.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    # Calculate key metrics
    portfolio_total = data.get('mainSuperBalance', 0) + data.get('sequencingBuffer', 0)
    annual_spending = data.get('baseSpending', 0)
    pension_income = data.get('totalPensionIncome', 0)
    net_drawdown = annual_spending - pension_income
    
    # Success/Fail Status
    mc_results = data.get('monteCarloResults') or data.get('historicalMonteCarloResults')
    if mc_results:
        success_rate = mc_results.get('successRate', 0)
        risk_text, risk_color = get_risk_level(success_rate)
        
        # Status banner
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = f"PORTFOLIO STATUS: {risk_text}"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(14)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, risk_color)
        
        doc.add_paragraph()
    
    # Key Metrics Table
    table = doc.add_table(rows=5 if mc_results else 4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    metrics = [
        ('Total Portfolio', format_currency(portfolio_total)),
        ('Annual Spending', format_currency(annual_spending)),
        ('Pension Income', format_currency(pension_income)),
        ('Net Annual Drawdown', format_currency(net_drawdown)),
    ]
    
    if mc_results:
        metrics.append(('Monte Carlo Success Rate', f"{success_rate:.1f}%"))
    
    for i, (label, value) in enumerate(metrics):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
        table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(30, 58, 138)
        table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Key Findings
    doc.add_heading("Key Findings", level=2)
    
    findings = []
    
    # Portfolio adequacy
    if portfolio_total > 0:
        withdrawal_rate = (net_drawdown / portfolio_total) * 100
        if withdrawal_rate <= 4:
            findings.append(f"✓ Sustainable withdrawal rate of {withdrawal_rate:.1f}% (recommended ≤4%)")
        elif withdrawal_rate <= 5:
            findings.append(f"⚠ Moderate withdrawal rate of {withdrawal_rate:.1f}% (slightly above 4% guideline)")
        else:
            findings.append(f"⚠ High withdrawal rate of {withdrawal_rate:.1f}% (exceeds 4% safe guideline)")
    
    # Success rate analysis
    if mc_results:
        if success_rate >= 85:
            findings.append(f"✓ Strong portfolio resilience with {success_rate:.1f}% success across 1,000 scenarios")
        elif success_rate >= 70:
            findings.append(f"⚠ Moderate portfolio resilience with {success_rate:.1f}% success rate")
        else:
            findings.append(f"⚠ Portfolio at risk with only {success_rate:.1f}% success rate")
    
    # Income sources
    if pension_income > 0:
        income_coverage = (pension_income / annual_spending) * 100
        findings.append(f"ℹ Pension income covers {income_coverage:.0f}% of spending needs")
    
    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')
    
    doc.add_page_break()

def create_assumptions_section(doc, data):
    """Create financial assumptions section"""
    
    heading = doc.add_heading("2. Financial Assumptions", level=1)
    heading.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    # Personal Details
    doc.add_heading("Personal Details", level=2)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    personal = [
        ('Current Age', str(data.get('currentAge', 'N/A'))),
        ('Retirement Age', str(data.get('retirementAge', 'N/A'))),
        ('Recipient Type', data.get('pensionRecipientType', 'single').title()),
        ('Homeowner Status', 'Yes' if data.get('isHomeowner') else 'No'),
    ]
    
    for i, (label, value) in enumerate(personal):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Financial Resources
    doc.add_heading("Financial Resources", level=2)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    resources = [
        ('Superannuation Balance', format_currency(data.get('mainSuperBalance', 0))),
        ('Sequencing Buffer', format_currency(data.get('sequencingBuffer', 0))),
        ('Total Portfolio', format_currency(data.get('mainSuperBalance', 0) + data.get('sequencingBuffer', 0))),
        ('Annual Pension Income', format_currency(data.get('totalPensionIncome', 0))),
    ]
    
    for i, (label, value) in enumerate(resources):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Bold and color the total row
        if i == 2:
            table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
            table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    doc.add_paragraph()
    
    # Spending & Economic Assumptions
    doc.add_heading("Spending & Economic Assumptions", level=2)
    
    assumptions = [
        ('Annual Base Spending', format_currency(data.get('baseSpending', 0))),
        ('Spending Pattern', data.get('spendingPattern', 'constant').title()),
        ('Expected Return', format_percent(data.get('selectedScenario', 0))),
        ('Inflation Rate', format_percent(data.get('inflationRate', 2.5))),
        ('Age Pension', 'Included' if data.get('includeAgePension') else 'Excluded'),
    ]
    
    # Add splurge if applicable
    if data.get('splurgeAmount', 0) > 0:
        assumptions.append((
            'Splurge Spending', 
            f"{format_currency(data.get('splurgeAmount'))} for {data.get('splurgeDuration')} years"
        ))
    
    table = doc.add_table(rows=len(assumptions), cols=2)
    table.style = 'Light Grid Accent 1'
    
    for i, (label, value) in enumerate(assumptions):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()

def create_projections_section(doc, data):
    """Create portfolio projections section"""
    
    heading = doc.add_heading("3. Portfolio Projections", level=1)
    heading.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    chart_data = data.get('chartData', [])
    if not chart_data:
        doc.add_paragraph("No projection data available.")
        doc.add_page_break()
        return
    
    # Condensed view for readability
    if len(chart_data) > 25:
        selected_years = (chart_data[:10] + 
                         chart_data[len(chart_data)//2-2:len(chart_data)//2+3] +
                         chart_data[-10:])
        condensed = True
    else:
        selected_years = chart_data
        condensed = False
    
    # Create table
    table = doc.add_table(rows=len(selected_years) + 1, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    headers = ['Year', 'Age', 'Total Balance', 'Income', 'Spending']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, '1E3A8A')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for idx, row_data in enumerate(selected_years, 1):
        balance = row_data.get('Total Balance', 0)
        income = row_data.get('Income', 0)
        spending = row_data.get('Spending', 0)
        
        values = [
            str(row_data.get('year', '')),
            str(row_data.get('age', '')),
            format_currency(balance) if isinstance(balance, (int, float)) else str(balance),
            format_currency(income) if isinstance(income, (int, float)) else str(income),
            format_currency(spending) if isinstance(spending, (int, float)) else str(spending),
        ]
        
        for i, value in enumerate(values):
            cell = table.rows[idx].cells[i]
            cell.text = value
            if i < 2:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    # Note
    if condensed:
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = note.add_run("Note: Table shows representative years. Export detailed CSV for complete year-by-year data.")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_page_break()

def create_risk_analysis_section(doc, data):
    """Create risk analysis section"""
    
    heading = doc.add_heading("4. Risk Analysis", level=1)
    heading.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    mc_results = data.get('monteCarloResults') or data.get('historicalMonteCarloResults')
    
    if not mc_results:
        doc.add_paragraph("No Monte Carlo analysis available. Run Monte Carlo simulation for comprehensive risk assessment.")
        doc.add_page_break()
        return
    
    success_rate = mc_results.get('successRate', 0)
    risk_text, risk_color = get_risk_level(success_rate)
    
    # Success Rate Banner
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = f"Monte Carlo Success Rate: {success_rate:.1f}%\n{risk_text}"
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(14)
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    set_cell_background(cell, risk_color)
    
    doc.add_paragraph()
    
    # Interpretation
    doc.add_heading("What This Means", level=2)
    
    interpretation = []
    if success_rate >= 90:
        interpretation.append("Excellent: Your portfolio is highly resilient across a wide range of market conditions.")
    elif success_rate >= 80:
        interpretation.append("Good: Your portfolio shows strong resilience in most market scenarios.")
    elif success_rate >= 70:
        interpretation.append("Moderate: Your portfolio succeeds in most scenarios but has meaningful risk in adverse conditions.")
    elif success_rate >= 60:
        interpretation.append("Caution: Your portfolio has significant risk of depletion in challenging markets.")
    else:
        interpretation.append("High Risk: Your portfolio is likely to deplete prematurely in many scenarios.")
    
    interpretation.append(f"Out of 1,000 simulated retirement scenarios with varying market returns, {int(success_rate*10)} scenarios maintained sufficient funds through retirement while {int((100-success_rate)*10)} scenarios depleted prematurely.")
    
    for text in interpretation:
        doc.add_paragraph(text)
    
    doc.add_paragraph()
    
    # Percentile Analysis
    percentiles = mc_results.get('percentiles', {})
    if percentiles:
        doc.add_heading("Portfolio Balance at Retirement End (Percentiles)", level=2)
        
        table = doc.add_table(rows=6, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header
        headers = ['Percentile', 'Final Balance', 'Interpretation']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            set_cell_background(cell, '475569')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Data
        p_items = [
            ('p10', '10th', 'Worst case (bottom 10%)'),
            ('p25', '25th', 'Below average'),
            ('p50', '50th', 'Median (typical outcome)'),
            ('p75', '75th', 'Above average'),
            ('p90', '90th', 'Best case (top 10%)'),
        ]
        
        for idx, (key, label, interp) in enumerate(p_items, 1):
            value = percentiles.get(key, 0)
            table.rows[idx].cells[0].text = label
            table.rows[idx].cells[1].text = format_currency(value)
            table.rows[idx].cells[2].text = interp
            table.rows[idx].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()

def create_recommendations_section(doc, data):
    """Create recommendations section"""
    
    heading = doc.add_heading("5. Recommendations", level=1)
    heading.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    recommendations = []
    
    # Calculate metrics
    portfolio_total = data.get('mainSuperBalance', 0) + data.get('sequencingBuffer', 0)
    annual_spending = data.get('baseSpending', 0)
    pension_income = data.get('totalPensionIncome', 0)
    net_drawdown = annual_spending - pension_income
    
    if portfolio_total > 0:
        withdrawal_rate = (net_drawdown / portfolio_total) * 100
    else:
        withdrawal_rate = 0
    
    mc_results = data.get('monteCarloResults') or data.get('historicalMonteCarloResults')
    success_rate = mc_results.get('successRate', 0) if mc_results else None
    
    # Build recommendations
    if withdrawal_rate > 5:
        recommendations.append({
            'priority': 'HIGH',
            'color': 'EF4444',
            'recommendation': f'Reduce withdrawal rate from {withdrawal_rate:.1f}% to 4-5%',
            'rationale': 'Current withdrawal rate significantly exceeds the 4% safe withdrawal guideline, increasing risk of portfolio depletion.',
            'action': f'Consider reducing annual spending by {format_currency(annual_spending - (portfolio_total * 0.04) - pension_income)} or increasing portfolio balance.'
        })
    elif withdrawal_rate > 4:
        recommendations.append({
            'priority': 'MEDIUM',
            'color': 'F59E0B',
            'recommendation': f'Monitor withdrawal rate of {withdrawal_rate:.1f}%',
            'rationale': 'Withdrawal rate is slightly above the conservative 4% guideline but within acceptable range.',
            'action': 'Consider modest spending reductions or portfolio growth strategies if markets underperform.'
        })
    
    if success_rate is not None:
        if success_rate < 70:
            recommendations.append({
                'priority': 'HIGH',
                'color': 'EF4444',
                'recommendation': f'Improve success rate from {success_rate:.1f}%',
                'rationale': 'Less than 70% success rate indicates significant risk of portfolio failure.',
                'action': 'Options: (1) Reduce spending, (2) Delay retirement, (3) Increase savings, (4) Review asset allocation for better risk-adjusted returns.'
            })
        elif success_rate < 85:
            recommendations.append({
                'priority': 'MEDIUM',
                'color': 'F59E0B',
                'recommendation': 'Strengthen portfolio resilience',
                'rationale': f'Success rate of {success_rate:.1f}% is adequate but could be improved for greater security.',
                'action': 'Consider modest adjustments to spending or contributions to reach 85%+ success rate target.'
            })
    
    recommendations.append({
        'priority': 'MEDIUM',
        'color': 'F59E0B',
        'recommendation': 'Maintain age-appropriate asset allocation',
        'rationale': 'Balanced portfolios typically target 60-70% growth assets (shares) with 30-40% defensive (bonds/cash) for retirees.',
        'action': 'Review your current allocation with a financial adviser to ensure it matches your risk tolerance and time horizon.'
    })
    
    buffer = data.get('sequencingBuffer', 0)
    if buffer < annual_spending * 2:
        recommendations.append({
            'priority': 'MEDIUM',
            'color': 'F59E0B',
            'recommendation': 'Increase sequencing buffer',
            'rationale': f'Current buffer of {format_currency(buffer)} provides less than 2 years of spending protection.',
            'action': f'Consider increasing buffer to {format_currency(annual_spending * 2)} (2 years of expenses) to protect against early market downturns.'
        })
    
    if not data.get('includeAgePension'):
        recommendations.append({
            'priority': 'LOW',
            'color': '10B981',
            'recommendation': 'Consider Age Pension eligibility',
            'rationale': 'Age Pension provides valuable inflation-indexed income that reduces portfolio drawdown.',
            'action': 'Check eligibility using Services Australia online estimator. Even partial pension entitlement can improve portfolio longevity.'
        })
    
    # Display recommendations
    for i, rec in enumerate(recommendations, 1):
        # Priority badge
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = f"{rec['priority']} PRIORITY"
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, rec['color'])
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.rows[0].height = Inches(0.3)
        
        # Recommendation content
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {rec['recommendation']}")
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(30, 58, 138)
        
        p = doc.add_paragraph()
        run = p.add_run("Rationale: ")
        run.font.italic = True
        run.font.bold = True
        run = p.add_run(rec['rationale'])
        
        p = doc.add_paragraph()
        run = p.add_run("Recommended Action: ")
        run.font.italic = True
        run.font.bold = True
        run = p.add_run(rec['action'])
        
        doc.add_paragraph()
    
    # Professional advice note
    p = doc.add_paragraph()
    run = p.add_run("Professional Advice: ")
    run.font.bold = True
    run = p.add_run("These recommendations are based on the assumptions and projections in this report. Consult with a licensed financial adviser to develop a personalized strategy that considers your complete financial situation, goals, and risk tolerance.")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    
    doc.add_page_break()

def create_scenario_details_section(doc, data):
    """Create scenario details section"""
    
    heading = doc.add_heading("6. Scenario Details", level=1)
    heading.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    
    # One-off expenses
    one_off = data.get('oneOffExpenses', [])
    if one_off and any(e.get('amount', 0) > 0 for e in one_off):
        doc.add_heading("Planned One-Off Expenses", level=2)
        
        # Filter expenses with amounts
        expenses = [e for e in one_off if e.get('amount', 0) > 0]
        
        table = doc.add_table(rows=len(expenses) + 1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header
        headers = ['Description', 'Age', 'Amount']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            set_cell_background(cell, '475569')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Data
        for idx, expense in enumerate(expenses, 1):
            table.rows[idx].cells[0].text = expense.get('description', 'N/A')
            table.rows[idx].cells[1].text = str(expense.get('age', 'N/A'))
            table.rows[idx].cells[2].text = format_currency(expense.get('amount', 0))
            table.rows[idx].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.rows[idx].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
    
    # Methodology
    doc.add_heading("Calculation Methodology", level=2)
    
    p = doc.add_paragraph("This analysis uses a year-by-year projection model that accounts for:")
    
    methodology_points = [
        "Investment returns (with optional stochastic variation via Monte Carlo)",
        "Inflation effects on spending and Age Pension",
        "Minimum superannuation drawdown requirements",
        "Australian Age Pension asset and income tests",
        "Dynamic spending adjustments based on selected pattern",
        "Sequencing risk buffer to protect against early market downturns",
    ]
    
    for point in methodology_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph("Monte Carlo simulations run 1,000 scenarios with randomized annual returns based on your expected return and volatility inputs, providing statistical confidence intervals for outcomes.")

def main():
    if len(sys.argv) != 3:
        print("Usage: generate_retirement_docx.py <input_json> <output_docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    # Load data
    with open(input_file, 'r') as f:
        data = json.load(f)
    
    # Create document
    doc = Document()
    
    # Set document properties
    core_props = doc.core_properties
    core_props.title = "Australian Retirement Planning Report"
    core_props.author = "Retirement Planning Calculator"
    core_props.subject = "Comprehensive Retirement Analysis"
    core_props.created = datetime.now()
    
    # Build document sections
    create_cover_page(doc, data)
    create_table_of_contents(doc)
    create_executive_summary(doc, data)
    create_assumptions_section(doc, data)
    create_projections_section(doc, data)
    create_risk_analysis_section(doc, data)
    create_recommendations_section(doc, data)
    create_scenario_details_section(doc, data)
    
    # Save document
    doc.save(output_file)
    
    print(f"Word document generated successfully: {output_file}")

if __name__ == "__main__":
    main()
